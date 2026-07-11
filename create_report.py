import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styling helper functions
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(24, 26, 34)
        p.paragraph_format.space_after = Pt(20)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = RGBColor(124, 77, 255)
        p.paragraph_format.space_after = Pt(40)

    def add_meta(label, value):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(1.5)
        p.paragraph_format.space_after = Pt(4)
        run_l = p.add_run(label + ": ")
        run_l.font.name = 'Arial'
        run_l.font.size = Pt(11)
        run_l.font.bold = True
        
        run_v = p.add_run(value)
        run_v.font.name = 'Calibri'
        run_v.font.size = Pt(11)

    def add_heading1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(24, 26, 34)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True

    def add_heading2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(124, 77, 255)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True

    def add_body(text, bold=False, italic=False, space_after=6):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = bold
        run.font.italic = italic
        p.paragraph_format.space_after = Pt(space_after)

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(40, 44, 56)

    # --- COVER PAGE ---
    add_title("\n\n\nSocialMusic")
    add_subtitle("An Interactive Music Player utilizing Custom DSA Structures")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(50) # Spacer
    
    add_meta("Submitted by", "Student 1, Student 2, Student 3, and Student 4")
    add_meta("Course/Program", "Bachelor of Technology in Computer Science & Engineering (B.Tech CSE)")
    add_meta("Department", "Department of Computer Science & Engineering")
    add_meta("Institution", "Lovely Professional University (LPU)")
    add_meta("Academic Session", "2026-2027")
    add_meta("Guide/Mentor Name", "LPU CSE Faculty Mentor")
    
    doc.add_page_break()

    # --- TABLE OF CONTENTS ---
    add_heading1("Table of Contents")
    add_body("1. Introduction............................................................................................................... Page 3")
    add_body("2. Problem Statement........................................................................................................ Page 3")
    add_body("3. Project Objectives......................................................................................................... Page 3")
    add_body("4. Project Modules & Team Roles........................................................................................ Page 4")
    add_body("5. Technologies Used......................................................................................................... Page 4")
    add_body("6. Module Description (Member-wise)................................................................................. Page 5")
    add_body("7. Data Structures & Algorithms Design................................................................................ Page 6")
    add_body("8. Future Scope & Conclusion............................................................................................... Page 7")
    add_body("9. References................................................................................................................ Page 7")
    
    doc.add_page_break()

    # --- CONTENT ---
    add_heading1("1. Introduction")
    add_body("SocialMusic is a custom Java Swing-based desktop application designed for playing music files while demonstrating advanced collaboration and fundamental Data Structures and Algorithms (DSA). The application includes dynamic WAV audio streaming, custom audio visualization, and a dark-themed user interface designed to feel modern, engaging, and premium.")

    add_heading1("2. Problem Statement")
    add_body("Typical audio players rely on pre-built collection libraries (like Java's LinkedList, ArrayList, PriorityQueue). While convenient, they abstract away low-level memory allocation, node linking, and array management operations. In an educational framework, building a music player from scratch allows students to learn, build, and optimize custom doubly linked lists, queue managers, adjacency lists, and search routines, coordinating them into a unified program.")

    add_heading1("3. Project Objectives")
    add_bullet("Develop a Doubly Linked List for active playlists, facilitating O(1) time complexity for track deletions and bidirectional traversal (Next/Previous).")
    add_bullet("Develop a First-In, First-Out (FIFO) Play Queue for queue-next capabilities with automated progression threads.")
    add_bullet("Develop a custom O(log N) Binary Search algorithm to query the library catalog alphabetically.")
    add_bullet("Develop a custom Recommendation Graph using adjacency lists to connection songs and provide dynamic recommendation links.")
    add_bullet("Ensure clean and robust object-oriented software design with modern flat dark-theme styling.")

    doc.add_page_break()

    add_heading1("4. Project Modules & Team Roles")
    add_body("The project was collaboratively designed by 4 students. Each student implemented their own data structure or algorithm module to guarantee complete autonomy and distinct contributions:")
    add_bullet("Student 1 (Playlist Management Module): Designed and built the custom PlaylistDLL and DLLNode structures, enabling bidirectional playlist operations.")
    add_bullet("Student 2 (Queue & Playback Automation Module): Programmed the FIFO PlayQueue and the playback loop integration, running background checks to automate transitioning from queues back to the playlist.")
    add_bullet("Student 3 (Catalog Search Module): Built the custom BinarySearch routine to perform fast lookups on titles.")
    add_bullet("Student 4 (Recommendation Module): Developed the custom RecommendationGraph structure using nodes and edges to model tracks and similarity.")

    add_heading1("5. Technologies Used")
    add_bullet("Programming Language: Java 21 Launcher")
    add_bullet("UI Framework: Java Swing & AWT (Abstract Window Toolkit)")
    add_bullet("Audio Engine: Java Sound Sampled API (javax.sound.sampled)")
    add_bullet("Version Control: Git & GitHub")

    doc.add_page_break()

    add_heading1("6. Module Description (Member-wise)")
    
    add_heading2("Module 1: Doubly Linked List Playlist (Student 1)")
    add_body("Playlist DLL implements a custom list layout where each song resides in a DLLNode pointing to both next and prev. When 'Next' or 'Previous' buttons are activated, pointer traversal occurs in O(1) time without rebuilding the track order array. Reordering operations ('Move Up' / 'Move Down') perform swift node data swaps.")
    add_body("UI/UX Integration Details:", bold=True)
    add_bullet("The DLL Playlist View: Renders the active playlist in a custom-styled dark-themed JTable.")
    add_bullet("Track Reordering: Provides 'Move Up' and 'Move Down' buttons that swap adjacent node values in the DLL. The JTable updates instantly to reflect the modified list order, providing direct visual feedback.")
    add_bullet("Play Button: Triggers the audio stream for the selected song in the playlist.")
    add_bullet("Remove Button: Unlinks the selected song from the list and refreshes the table row layout.")
    add_bullet("Quick Add Selector: A JComboBox allows users to select any song from the global catalog and append it to the DLL.")
    add_bullet("Sidebar Notification Badge: Dynamically displays the size of the playlist as DLL Playlist (N) in real-time.")

    add_heading2("Module 2: FIFO Play Queue (Student 2)")
    add_body("Stages pending songs using a custom Queue structure. Enqueuing appends a song at the tail, and dequeuing retrieves it from the head. The custom playback automation runs upon a track finishing: it polls the play queue first. If a song exists, it plays it; otherwise, it advances the active DLL playlist.")
    add_body("UI/UX Integration Details:", bold=True)
    add_bullet("The FIFO Queue View: Shows upcoming tracks in their exact play order.")
    add_bullet("Staging Controls: Includes an 'Add to Queue' button on the Dashboard and a 'Quick Add' dropdown selector inside the Queue panel to push tracks onto the tail of the play queue.")
    add_bullet("Playback Automation: Automatically polls the queue when a track ends. If populated, it dequeues the head song and plays it immediately. If empty, it falls back to the active Playlist DLL to continue playing.")
    add_bullet("Sidebar Badge: Real-time counter badge FIFO Queue (N) updates whenever a song is enqueued or dequeued.")

    doc.add_page_break()

    add_heading2("Module 3: Binary Search Catalog (Student 3)")
    add_body("Allows logarithmic search over the library. The catalog is kept sorted, and binary search repeatedly halves the search interval until the query title matches or the boundary collapses, completing searching in O(log N) runtime.")
    add_body("Detailed Algorithmic Mechanics:", bold=True)
    add_bullet("Binary Search uses a divide-and-conquer strategy to locate a target item in logarithmic O(log N) time.")
    add_bullet("The Sorting Precondition: Because Binary Search splits the search space in half by comparing with a midpoint, it requires the search space to be sorted.")
    add_bullet("The Search Flow:")
    add_body("1. When the user types a search query in txtSearch and clicks the 'Search' button, the event handler first sorts the catalog list alphabetically by title:")
    add_code("catalog.sort((a, b) -> a.title.compareToIgnoreCase(b.title));")
    add_body("2. The Dashboard JList updates to reflect this sorted order.")
    add_body("3. The binary search starts lookups. It computes the midpoint mid = (low + high) / 2.")
    add_body("4. If the song at the mid index matches the query (case-insensitive), it returns mid.")
    add_body("5. If the target query is lexicographically larger than the midpoint song title, low is moved to mid + 1. Otherwise, high is moved to mid - 1.")
    add_body("6. If no match is found, it returns -1.")
    add_body("UI/UX Selection & Scrolling Details:", bold=True)
    add_bullet("Highlight Selection: If a match is found, the JList selection is updated to highlight the matching row: list.setSelectedIndex(idx).")
    add_bullet("Auto-Scroll: It automatically scrolls the list viewport to make the highlighted selection visible: list.ensureIndexIsVisible(idx).")
    add_bullet("Update Details Panel: The selected song details card (Title, Artist, Genre) updates immediately.")
    add_bullet("Success Dialogue: A dialog box pops up confirming: 'Found: [Song Title]'.")
    add_bullet("Immediate Action: Once highlighted, the user can click 'PLAY NOW' to start the music or add it to the playlist/queue.")

    add_heading2("Module 4: Recommendation Graph (Student 4)")
    add_body("Leverages an adjacency list representation where each song is a node, and edges signify recommendation links. It maps relations, and returns the neighbor nodes in O(1) lookup time to populate the Recommendations list.")
    add_body("UI/UX Integration Details:", bold=True)
    add_bullet("The Recommendations Tab: Displays dynamic recommendations based on connections in the graph.")
    add_bullet("Dynamic Synchronization: Selecting a song on the Dashboard or playing a track triggers updateRecommendations(...) to load all adjacent neighbor nodes of that track.")
    add_bullet("Immediate Exploration: Clicking 'Play Selected Recommendation' starts playback of the suggested track and recalculates recommendations for the new song on the fly.")

    doc.add_page_break()

    add_heading1("7. Data Structures & Algorithms Design")
    
    add_heading2("Doubly Linked List Structure")
    add_code("class DLLNode {")
    add_code("    Song song;")
    add_code("    DLLNode prev, next;")
    add_code("}")

    add_heading2("FIFO Queue Structure")
    add_code("class QueueNode {")
    add_code("    Song song;")
    add_code("    QueueNode next;")
    add_code("}")

    add_heading2("Binary Search Algorithm")
    add_code("public static int searchByTitle(List<Song> list, String query) {")
    add_code("    int low = 0, high = list.size() - 1;")
    add_code("    while (low <= high) {")
    add_code("        int mid = (low + high) / 2;")
    add_code("        int cmp = list.get(mid).title.compareToIgnoreCase(query);")
    add_code("        if (cmp == 0) return mid;")
    add_code("        else if (cmp < 0) low = mid + 1;")
    add_code("        else high = mid - 1;")
    add_code("    }")
    add_code("    return -1;")
    add_code("}")

    add_heading2("Recommendation Graph Structure")
    add_code("class GraphNode {")
    add_code("    Song song;")
    add_code("    List<Song> neighbors = new ArrayList<>();")
    add_code("}")

    doc.add_page_break()

    add_heading1("8. Future Scope & Conclusion")
    add_body("The current application provides a robust showcase of core DSA concepts in a clean Swing GUI. Future enhancements include:")
    add_bullet("Persistent SQL backend database integration.")
    add_bullet("Socket-based network playlists to support live collaborative streaming.")
    add_bullet("Digital Signal Processing (DSP) equalizer controls.")
    add_body("Conclusion: The project successfully integrates four distinct student-built structures into a cohesive application, showing the power of modular design and standard algorithms.")

    add_heading1("9. References")
    add_bullet("Cormen, T. H. (2009). Introduction to Algorithms (3rd ed.). MIT Press.")
    add_bullet("Oracle Java Sound Documentation. Tutorial: Sampled Audio API.")
    add_bullet("GeeksforGeeks Data Structures & Algorithms. Reference Materials.")

    doc.save("Project_Report.docx")
    print("Report generated successfully as Project_Report.docx")

if __name__ == "__main__":
    create_report()
